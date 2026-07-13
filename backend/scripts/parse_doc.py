"""PDF/DOCX/CSV/JSON/XLSX/XLS 文件解析工具
Usage: python parse_doc.py <file_path> [type: pdf|docx|csv|json|xlsx|xls]
Outputs parsed text to stdout (UTF-8)
"""
import sys, json, os

def parse_pdf(path):
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        pages = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                pages.append(text.strip())
        doc.close()
        return '\n\n'.join(pages)
    except ImportError:
        return '[PDF解析失败: 请安装PyMuPDF模块 (pip install PyMuPDF)]'
    except Exception as e:
        return f'[PDF解析失败: {str(e)}]'

def parse_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    paragraphs.append(row_text)
        return '\n\n'.join(paragraphs)
    except ImportError:
        return '[DOCX解析失败: 请安装python-docx模块 (pip install python-docx)]'
    except Exception as e:
        return f'[DOCX解析失败: {str(e)}]'

def parse_csv(path):
    try:
        with open(path, 'r', encoding='utf-8-sig') as f:
            content = f.read()
        return content
    except Exception as e:
        return f'[CSV解析失败: {str(e)}]'

def parse_xlsx(path):
    """解析 Excel 文件 (.xlsx / .xls)，提取所有工作表的文本内容"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        all_sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_list = []
            for row in ws.iter_rows(values_only=True):
                # 跳过全空行
                row_values = [str(cell) if cell is not None else '' for cell in row]
                if any(v.strip() for v in row_values):
                    rows_list.append(' | '.join(row_values))
            if rows_list:
                all_sheets.append(f'=== Sheet: {sheet_name} ===\n' + '\n'.join(rows_list))
            else:
                all_sheets.append(f'=== Sheet: {sheet_name} ===\n(空)')
        wb.close()
        return '\n\n'.join(all_sheets)
    except ImportError:
        return '[Excel解析失败: 请安装openpyxl模块 (pip install openpyxl)]'
    except Exception as e:
        return f'[Excel解析失败: {str(e)}]'

def parse_json_file(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
    except Exception as e:
        return f'[JSON解析失败: {str(e)}]'

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('[错误: 请提供文件路径]')
        sys.exit(1)

    file_path = sys.argv[1]
    file_type = sys.argv[2].lower() if len(sys.argv) > 2 else os.path.splitext(file_path)[1].lstrip('.').lower()

    if not os.path.exists(file_path):
        print(f'[错误: 文件不存在: {file_path}]')
        sys.exit(1)

    parsers = {'pdf': parse_pdf, 'docx': parse_docx, 'csv': parse_csv, 'json': parse_json_file, 'xlsx': parse_xlsx, 'xls': parse_xlsx}
    parser = parsers.get(file_type)
    if not parser:
        print(f'[错误: 不支持的文件类型: {file_type}]')
        sys.exit(1)

    result = parser(file_path)
    print(result)
